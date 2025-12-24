"""
Research Paper Validator
Validates whether an uploaded PDF or DOCX is a legitimate research paper.
"""

import re
from typing import Tuple, List
import fitz  # PyMuPDF
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ResearchPaperValidator:
    """
    Validates if a PDF or DOCX document is a research paper based on:
    - File format (PDF or DOCX)
    - Page count (>= 4 pages)
    - Section presence (Abstract, Introduction, Methodology, Results, Conclusion, References)
    - Citation patterns ([1], et al., DOI)
    """
    
    # Standard research paper sections
    REQUIRED_SECTIONS = [
        "abstract",
        "introduction",
        "methodology",
        "method",
        "results",
        "conclusion",
        "references",
        "bibliography",
        "related work",
        "literature review"
    ]
    
    # Minimum number of sections required
    MIN_SECTIONS_REQUIRED = 3
    
    # Minimum page count
    MIN_PAGE_COUNT = 4
    
    def __init__(self):
        # Citation patterns
        self.citation_patterns = [
            r'\[\d+\]',                    # [1], [23]
            r'\b\w+\s+et\s+al\.',          # "Smith et al."
            r'doi:\s*10\.\d+',             # DOI patterns
            r'\(\d{4}\)',                  # Year citations like (2023)
            r'\w+\s+\(\d{4}\)',            # Author (Year) format
        ]
    
    def validate(self, file_path: str) -> Tuple[bool, str]:
        """
        Validate if the PDF or DOCX is a research paper.
        
        Args:
            file_path: Path to the PDF or DOCX file
            
        Returns:
            Tuple of (is_valid, error_message)
            If valid: (True, "")
            If invalid: (False, "Error description")
        """
        try:
            if file_path.lower().endswith('.docx'):
                return self._validate_docx(file_path)
            else:
                return self._validate_pdf(file_path)
        except Exception as e:
            return False, f"Error validating file: {str(e)}"
    
    def _validate_pdf(self, pdf_path: str) -> Tuple[bool, str]:
        """Validate PDF file."""
        # Open PDF
        doc = fitz.open(pdf_path)
        
        # Check 1: Minimum page count
        page_count = len(doc)
        if page_count < self.MIN_PAGE_COUNT:
            doc.close()
            return False, f"Only research papers are allowed. (Reason: Document has {page_count} pages, minimum {self.MIN_PAGE_COUNT} required)"
        
        # Extract text from all pages
        full_text = ""
        for page in doc:
            full_text += page.get_text()
        
        doc.close()
        
        # Check 2: Detect sections
        sections_found = self._detect_sections(full_text)
        if len(sections_found) < self.MIN_SECTIONS_REQUIRED:
            return False, f"Only research papers are allowed. (Reason: Found {len(sections_found)} sections, minimum {self.MIN_SECTIONS_REQUIRED} required)"
        
        # Check 3: Detect citations
        has_citations = self._detect_citations(full_text)
        if not has_citations:
            return False, "Only research papers are allowed. (Reason: No citation patterns detected)"
        
        # All checks passed
        return True, ""
    
    def _validate_docx(self, docx_path: str) -> Tuple[bool, str]:
        """Validate DOCX file."""
        if not DOCX_AVAILABLE:
            return False, "python-docx is required for DOCX support"
        
        doc = Document(docx_path)
        
        # Count paragraphs as proxy for pages (approx 45 paragraphs per page)
        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
        estimated_pages = max(1, paragraph_count // 45)
        
        # Check 1: Minimum page count
        if estimated_pages < self.MIN_PAGE_COUNT:
            return False, f"Only research papers are allowed. (Reason: Document has approximately {estimated_pages} pages, minimum {self.MIN_PAGE_COUNT} required)"
        
        # Extract text
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Check 2: Detect sections
        sections_found = self._detect_sections(full_text)
        if len(sections_found) < self.MIN_SECTIONS_REQUIRED:
            return False, f"Only research papers are allowed. (Reason: Found {len(sections_found)} sections, minimum {self.MIN_SECTIONS_REQUIRED} required)"
        
        # Check 3: Detect citations
        has_citations = self._detect_citations(full_text)
        if not has_citations:
            return False, "Only research papers are allowed. (Reason: No citation patterns detected)"
        
        # All checks passed
        return True, ""
    
    def _detect_sections(self, text: str) -> List[str]:
        """
        Detect standard research paper sections in the text.
        
        Args:
            text: Full text of the document
            
        Returns:
            List of detected section names
        """
        text_lower = text.lower()
        detected_sections = []
        
        for section in self.REQUIRED_SECTIONS:
            # Look for section headers (e.g., "1. Introduction", "Abstract", "METHODOLOGY")
            # Pattern matches section name at start of line or after number
            patterns = [
                rf'\b{section}\b',                    # Basic match
                rf'^\s*\d+\.?\s*{section}',           # "1. Introduction"
                rf'^\s*{section}\s*$',                # Section on its own line
                rf'\n\s*{section}\s*\n',              # Section between newlines
            ]
            
            for pattern in patterns:
                if re.search(pattern, text_lower, re.MULTILINE | re.IGNORECASE):
                    if section not in detected_sections:
                        detected_sections.append(section)
                    break
        
        return detected_sections
    
    def _detect_citations(self, text: str) -> bool:
        """
        Detect citation patterns in the text.
        
        Args:
            text: Full text of the document
            
        Returns:
            True if citations are found, False otherwise
        """
        citation_count = 0
        
        for pattern in self.citation_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            citation_count += len(matches)
            
            # If we find at least 3 citations, it's likely a research paper
            if citation_count >= 3:
                return True
        
        return citation_count >= 3
    
    def get_document_info(self, pdf_path: str) -> dict:
        """
        Extract basic information about the PDF document.
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Dictionary with document information
        """
        try:
            doc = fitz.open(pdf_path)
            
            # Extract metadata
            metadata = doc.metadata
            page_count = len(doc)
            
            # Get first page text for title extraction
            first_page_text = doc[0].get_text() if page_count > 0 else ""
            
            doc.close()
            
            return {
                "page_count": page_count,
                "title": metadata.get("title", "Unknown"),
                "author": metadata.get("author", "Unknown"),
                "first_page_preview": first_page_text[:500]
            }
        except Exception as e:
            return {"error": str(e)}

