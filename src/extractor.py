"""
Document Text Extractor
Extracts text from research papers (PDF and DOCX) and identifies sections.
"""

import re
from typing import List, Dict, Tuple
import fitz  # PyMuPDF
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class PDFExtractor:
    """
    Extracts text from PDF and DOCX research papers and identifies sections.
    """
    
    # Common section headers in research papers
    SECTION_HEADERS = [
        "abstract",
        "introduction",
        "background",
        "related work",
        "literature review",
        "methodology",
        "method",
        "methods",
        "approach",
        "experimental setup",
        "experiment",
        "results",
        "findings",
        "discussion",
        "results and discussion",
        "evaluation",
        "conclusion",
        "conclusions",
        "future work",
        "references",
        "bibliography",
        "acknowledgments",
        "appendix"
    ]
    
    def __init__(self):
        self.section_pattern = self._build_section_pattern()
    
    def _build_section_pattern(self) -> re.Pattern:
        """
        Build a regex pattern to match section headers.
        Matches patterns like:
        - "1. Introduction"
        - "METHODOLOGY"
        - "2.1 Experimental Setup"
        """
        sections_regex = "|".join(self.SECTION_HEADERS)
        pattern = rf'^[\s]*(\d+\.?\d*\.?\s*)?({sections_regex})[\s]*$'
        return re.compile(pattern, re.IGNORECASE | re.MULTILINE)
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract all text from a PDF or DOCX.
        
        Args:
            file_path: Path to the PDF or DOCX file
            
        Returns:
            Full text content
        """
        if file_path.lower().endswith('.docx'):
            return self._extract_text_docx(file_path)
        else:
            return self._extract_text_pdf(file_path)
    
    def _extract_text_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF."""
        doc = fitz.open(pdf_path)
        full_text = ""
        
        for page in doc:
            full_text += page.get_text()
        
        doc.close()
        return full_text
    
    def _extract_text_docx(self, docx_path: str) -> str:
        """Extract text from DOCX."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX support")
        
        doc = Document(docx_path)
        full_text = []
        
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        return '\n'.join(full_text)
    
    def extract_with_page_numbers(self, file_path: str) -> List[Dict[str, any]]:
        """
        Extract text with page number information.
        
        Args:
            file_path: Path to the PDF or DOCX file
            
        Returns:
            List of dictionaries with page number and text
        """
        if file_path.lower().endswith('.docx'):
            return self._extract_with_page_numbers_docx(file_path)
        else:
            return self._extract_with_page_numbers_pdf(file_path)
    
    def _extract_with_page_numbers_pdf(self, pdf_path: str) -> List[Dict[str, any]]:
        """Extract text with page numbers from PDF."""
        doc = fitz.open(pdf_path)
        pages_data = []
        
        for page_num, page in enumerate(doc, start=1):
            pages_data.append({
                "page_number": page_num,
                "text": page.get_text()
            })
        
        doc.close()
        return pages_data
    
    def _extract_with_page_numbers_docx(self, docx_path: str) -> List[Dict[str, any]]:
        """Extract text with page numbers from DOCX (simulated pages)."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX support")
        
        doc = Document(docx_path)
        
        # DOCX doesn't have page breaks in the API, so we simulate pages
        # by grouping paragraphs (approx 40-50 paragraphs per page)
        paragraphs_per_page = 45
        pages_data = []
        current_page = []
        
        for i, para in enumerate(doc.paragraphs):
            current_page.append(para.text)
            
            if (i + 1) % paragraphs_per_page == 0:
                pages_data.append({
                    "page_number": len(pages_data) + 1,
                    "text": '\n'.join(current_page)
                })
                current_page = []
        
        # Add remaining paragraphs as last page
        if current_page:
            pages_data.append({
                "page_number": len(pages_data) + 1,
                "text": '\n'.join(current_page)
            })
        
        return pages_data
    
    def extract_sections(self, file_path: str) -> List[Dict[str, any]]:
        """
        Extract text organized by detected sections.
        
        Args:
            file_path: Path to the PDF or DOCX file
            
        Returns:
            List of dictionaries containing section information
        """
        pages_data = self.extract_with_page_numbers(file_path)
        full_text = "\n".join([p["text"] for p in pages_data])
        
        # Split text into lines for section detection
        lines = full_text.split('\n')
        
        sections = []
        current_section = {
            "title": "Header",  # Before first section
            "content": [],
            "start_line": 0
        }
        
        for line_num, line in enumerate(lines):
            # Check if line is a section header
            match = self.section_pattern.match(line.strip())
            
            if match and len(line.strip()) < 100:  # Section headers are typically short
                # Save previous section
                if current_section["content"]:
                    current_section["content"] = '\n'.join(current_section["content"])
                    sections.append(current_section)
                
                # Start new section
                section_title = line.strip()
                current_section = {
                    "title": section_title,
                    "content": [],
                    "start_line": line_num
                }
            else:
                # Add line to current section
                if line.strip():  # Skip empty lines
                    current_section["content"].append(line)
        
        # Add the last section
        if current_section["content"]:
            current_section["content"] = '\n'.join(current_section["content"])
            sections.append(current_section)
        
        return sections
    
    def extract_metadata(self, file_path: str) -> Dict[str, str]:
        """
        Extract metadata from PDF or DOCX.
        
        Args:
            file_path: Path to the PDF or DOCX file
            
        Returns:
            Dictionary with metadata
        """
        if file_path.lower().endswith('.docx'):
            return self._extract_metadata_docx(file_path)
        else:
            return self._extract_metadata_pdf(file_path)
    
    def _extract_metadata_pdf(self, pdf_path: str) -> Dict[str, str]:
        """Extract metadata from PDF."""
        doc = fitz.open(pdf_path)
        metadata = doc.metadata
        
        # Try to extract title from first page if not in metadata
        title = metadata.get("title", "")
        if not title:
            first_page = doc[0].get_text()
            # Assume first non-empty line might be the title
            lines = [l.strip() for l in first_page.split('\n') if l.strip()]
            title = lines[0] if lines else "Untitled"
        
        doc.close()
        
        return {
            "title": title,
            "author": metadata.get("author", "Unknown"),
            "subject": metadata.get("subject", ""),
            "keywords": metadata.get("keywords", ""),
            "creator": metadata.get("creator", ""),
            "producer": metadata.get("producer", ""),
            "creation_date": metadata.get("creationDate", ""),
        }
    
    def _extract_metadata_docx(self, docx_path: str) -> Dict[str, str]:
        """Extract metadata from DOCX."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX support")
        
        doc = Document(docx_path)
        core_props = doc.core_properties
        
        # Try to extract title from first paragraph if not in metadata
        title = core_props.title or ""
        if not title and doc.paragraphs:
            # Use first non-empty paragraph as title
            for para in doc.paragraphs:
                if para.text.strip():
                    title = para.text.strip()
                    break
        
        return {
            "title": title or "Untitled",
            "author": core_props.author or "Unknown",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "creator": core_props.author or "",
            "producer": "",
            "creation_date": str(core_props.created) if core_props.created else "",
        }
    
    def extract_abstract(self, file_path: str) -> str:
        """
        Extract the abstract section from the paper.
        
        Args:
            file_path: Path to the PDF or DOCX file
            
        Returns:
            Abstract text or empty string if not found
        """
        sections = self.extract_sections(file_path)
        
        for section in sections:
            if "abstract" in section["title"].lower():
                return section["content"]
        
        return ""
